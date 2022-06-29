# -*- coding: utf-8 -*-
"""
Created on Thu Feb 18 13:15:44 2021

@author: naroa

This program saves information from a docx syllabus into a csv database.

"""
# import the modules
from docx import Document
import csv as csv 

# read the syllabus
print("Reading information from syllabus docx...")
document = Document('syllabi/female1.docx')

# extract information from syllabus
course_code = document.paragraphs[0].text.split('\t')[0]
course_name = document.paragraphs[0].text.split('\t')[1].lower()
semester_year = document.paragraphs[1].text.split('\t')[1] 
semester = semester_year.split(' ')[0]
year = semester_year.split(' ')[1]
whole_name = document.paragraphs[2].text.split('\t')[1]
first_name = whole_name.split(' ')[0]
last_name = whole_name.split(' ')[1]
telephone = document.paragraphs[3].text.split('\t')[1]
email = document.paragraphs[4].text.split('\t')[2]
office = document.paragraphs[5].text.split('\t')[2]
office_hours = document.paragraphs[6].text.split('\t')[1]
lectures = document.paragraphs[9].text.split(': ')[1]
textbook = document.paragraphs[11].text
attendance = document.paragraphs[13].text.split(' ')[3]
if attendance == ('not'):
    attendance = 'NOT REQUIRED'
else:
    attendance = 'REQUIRED'

try:
    assig_count = int(document.paragraphs[14].text.split(' ')[4])
except Exception:
    assig_count = 0

if assig_count != 0:
    assig_dates = []
    for i in range(1,assig_count+1):
        assig_dates.append(document.paragraphs[14+i].text.split('\t')[1])
else:
    assig_dates = 'There are no assignments in this course.'      

i=0 
try:    
    present_count = int(document.paragraphs[14+assig_count+1].text.split(' ')[4])
except Exception:
    present_count = 0

if present_count != 0:
    present_dates = []
    for i in range(1,present_count+1):
        present_dates.append(document.paragraphs[14+assig_count+1+i].text.split('\t')[1])
else:
    present_dates = 'There are no presentations in this course.'
   
midterm_date = document.paragraphs[14+assig_count+1+i+2].text.split('\t')[1]
final_date = document.paragraphs[14+assig_count+1+i+3].text.split('\t')[1]

if assig_count != 0: 
    assig_grade = document.paragraphs[14+assig_count+1+i+7].text.split('\t')[1]
else:
    assig_grade = 'There are no assignments in this course.' 

if present_count != 0:
    present_grade = document.paragraphs[14+assig_count+1+i+8].text.split('\t')[1]
else:
    present_grade = 'There are no presentations in this course.'

midterm_grade = document.paragraphs[-3].text.split('\t')[1]
final_grade = document.paragraphs[-2].text.split('\t')[1]

print(assig_grade)
print(present_grade)
print(midterm_grade)
print(final_grade)

late_submission = document.paragraphs[-1].text.split(' ')[8]

# save information in csv database
fields=[course_code,course_name,semester,year,first_name,last_name,telephone,email,office,office_hours,lectures,textbook,attendance,assig_dates,present_dates,midterm_date,final_date,assig_grade,present_grade,midterm_grade,final_grade,late_submission]
with open('course_db/courses_table.csv', 'a',newline='',encoding='utf-8') as f:
    writer = csv.writer(f,delimiter=',')
    writer.writerow(fields) 
    print("Saving information into csv database...")
f.close()

print("Information saved correctly.")
